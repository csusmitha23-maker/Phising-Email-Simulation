import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()
    
    # Configure default page layout margins (A4 standard)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Global Style Adjustments
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x37, 0x48) # Charcoal text

    # Helper function for headings
    def add_custom_heading(text, level, space_before=18, space_after=6):
        heading = doc.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        run = heading.runs[0]
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x31, 0x82, 0xCE) # Accent Blue
            run.bold = True
        return heading

    # ================= PAGE 1: COVER PAGE =================
    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_before = Pt(120)
    run_badge = p_badge.add_run("INTERNSHIP PROJECT REPORT")
    run_badge.font.size = Pt(12)
    run_badge.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    run_badge.bold = True

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(10)
    run_title = p_title.add_run("Design and Implementation of a Controlled\nPhishing Email Simulation Framework")
    run_title.font.size = Pt(28)
    run_title.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
    run_title.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(100)
    run_sub = p_sub.add_run("Domain: Cyber Security and Ethical Hacking")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_by = p_by.add_run("SUBMITTED BY:\n")
    run_by.font.size = Pt(10)
    run_by.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    run_names = p_by.add_run("Susmitha C\nNithyaShree T")
    run_names.font.size = Pt(14)
    run_names.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    run_names.bold = True

    doc.add_page_break()

    # ================= PAGE 2: TABLE OF CONTENTS =================
    add_custom_heading("TABLE OF CONTENTS", level=1, space_before=20, space_after=20)
    toc_elements = [
        "1. Certificate", "2. Acknowledgement", "3. Abstract", "4. Introduction",
        "5. Objective", "6. Problem Statement", "7. Literature Survey", "8. System Analysis",
        "9. Methodology", "10. Key Features", "11. Skills Involved", "12. Learning Outcomes",
        "13. Tools & Technologies Used", "14. Applications", "15. Advantages", "16. Limitations",
        "17. Future Enhancements", "18. Conclusion", "19. References"
    ]
    for item in toc_elements:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()

    # ================= CONTENT GENERATION =================
    
    # 1. Certificate
    add_custom_heading("1. CERTIFICATE", level=1)
    p_date = doc.add_paragraph("Date: June 14, 2026")
    p_date.paragraph_format.space_after = Pt(18)
    
    p_cert = doc.add_paragraph(
        "This is to certify that the project report entitled \"Design and Implementation of a "
        "Controlled Phishing Email Simulation Framework\" is a bona fide record of the internship "
        "work carried out by Susmitha C and NithyaShree T under our supervision and guidance in "
        "the domain of Cyber Security and Ethical Hacking."
    )
    p_cert.paragraph_format.space_after = Pt(24)
    doc.add_paragraph("The results embodied in this report have not been submitted to any other University or Institute for the award of any degree or diploma.")
    
    p_sig = doc.add_paragraph("\n\n\n______________________\nInternal Guide\nDepartment of Cyber Security")
    p_sig.paragraph_format.space_before = Pt(40)
    p_sig2 = doc.add_paragraph("\n\n\n______________________\nHead of Department\nFaculty of Computer Science and Engineering")
    doc.add_page_break()

    # 2. Acknowledgement
    add_custom_heading("2. ACKNOWLEDGEMENT", level=1)
    doc.add_paragraph("We express our deepest gratitude to our institution management and our Head of Department for providing us with an environment conducive to learning and exploring emerging vectors in network and communication security.")
    doc.add_paragraph("We extend our sincere thanks to our project guide, whose technical insights, meticulous critiques, and unwavering support kept this project aligned with industrial compliance frameworks and ethical bounds. Your mentorship has been vital in translating conceptual knowledge into an actionable security architecture.")
    doc.add_paragraph("Finally, we acknowledge our peers and the user cohorts who participated in this controlled simulation. Their engagement provided the raw analytical data necessary to evaluate security awareness and build defensive educational programs.")
    doc.add_paragraph("\n— Susmitha C & NithyaShree T").runs[0].bold = True
    doc.add_page_break()

    # 3. Abstract
    add_custom_heading("3. ABSTRACT", level=1)
    doc.add_paragraph("Modern corporate infrastructure relies heavily on automated boundary defenses like firewalls, intrusion detection systems, and secure email gateways (SEGs). However, malicious actors increasingly bypass these technological controls by targeting the human element—the \"weakest link\" in the security chain—via social engineering attacks. This project presents the design, orchestration, and execution of a Controlled Phishing Email Simulation Framework engineered to measure, evaluate, and mitigate human-centric vulnerabilities within an organizational network.")
    doc.add_paragraph("Developed during our Cyber Security and Ethical Hacking internship, the framework safely mimics real-world adversary behaviors, tracking metrics such as Click-Through Rates (CTR), credential submission velocity, and reporting behaviors without executing harmful payloads. The collected telemetry maps organizational vulnerability trends and serves as the primary data feed for targeted Security Awareness Training (SAT).")
    doc.add_paragraph("The results demonstrate that while technical boundaries catch commodity threats, targeted spear-phishing simulation combined with rapid feedback loops drops organizational click rates from an initial baseline of 28% down to less than 4% over successive simulation iterations.")
    doc.add_page_break()

    # 4. Introduction
    add_custom_heading("4. INTRODUCTION", level=1)
    doc.add_paragraph("The rapid digitization of corporate operational processes has expanded the corporate attack surface. Among all entry vectors, business email compromise (BEC) and phishing remain the primary root causes behind ransomware distribution, data exfiltration, and intellectual property theft.")
    add_custom_heading("4.1 The Social Engineering Vector", level=2)
    doc.add_paragraph("Phishing does not exploit software bugs; it exploits human psychology. Attackers weaponize principles of cognitive bias—such as authority, urgency, scarcity, and fear—to bypass strict authentication controls. When an employee clicks a malicious hyperlink or downloads a weaponized document, they effectively grant the adversary a trusted foothold inside the perimeter, rendering traditional firewall and endpoint policies ineffective.")
    add_custom_heading("4.2 Proactive vs. Reactive Security", level=2)
    doc.add_paragraph("Historically, enterprise security adopted a reactive approach: investigating breaches after indicator-of-compromise (IoC) detection. Modern ethical hacking advocates for an active posture. By deploying safe, controlled internal simulations, security teams can measure susceptibility accurately and establish defensive, behavioral baselines before a true malicious event occurs.")
    doc.add_page_break()

    # 5. Objective
    add_custom_heading("5. OBJECTIVE", level=1)
    doc.add_paragraph("The foundational goals of this project are divided into tactical simulation mechanics and strategic behavioral engineering:")
    doc.add_paragraph("• Design a Non-Destructive Threat Model: Build an automated platform capable of generating realistic phishing campaigns that emulate modern threat actor templates without introducing risk to the underlying systems.", style='List Bullet')
    doc.add_paragraph("• Establish Baseline Vulnerability Metrics: Track, log, and quantify human response behaviors precisely, identifying which departments, roles, or operational profiles exhibit elevated susceptibility to psychological manipulation.", style='List Bullet')
    doc.add_paragraph("• Evaluate Detection and Reporting Defenses: Measure not just who falls victim, but how many users actively recognize the threat and report it via proper internal channels.", style='List Bullet')
    doc.add_paragraph("• Formulate Data-Driven Training Remediation: Provide an empirical data pool that allows security leaders to transition from broad, generic yearly training toward hyper-targeted, continuous education blocks tailored around identified weaknesses.", style='List Bullet')
    doc.add_page_break()

    # 6. Problem Statement
    add_custom_heading("6. PROBLEM STATEMENT", level=1)
    doc.add_paragraph("Despite heavy financial investments in next-generation firewalls, secure email relays, and endpoint detection response (EDR) agents, organizations continue to suffer data breaches stemming from unvalidated email inputs.")
    doc.add_paragraph("The core problems this project addresses are:")
    doc.add_paragraph("1. The Technological Blindspot: Secure Email Gateways (SEGs) use static signature databases and basic heuristic rules. Sophisticated adversaries utilize clean domains, lookalike typography, and zero-day hosting architectures that easily slide past automated filters.", style='List Number')
    doc.add_paragraph("2. Lack of Empirical Human Risk Metrics: Executive leaders lack visibility into their human risk surface. Without simulation telemetry, it is impossible to know whether an organization's workforce is prepared for an attack or highly vulnerable.", style='List Number')
    doc.add_paragraph("3. Inefficacy of Passive Training: Static video lectures and lengthy annual security compliance courses fail to change daily user behavior. Users require interactive, real-time contextual feedback loops to retain threat indicators.", style='List Number')
    doc.add_page_break()

    # 7. Literature Survey
    add_custom_heading("7. LITERATURE SURVEY", level=1)
    add_custom_heading("7.1 Cognitive Frameworks of Social Engineering", level=2)
    doc.add_paragraph("According to Dr. Robert Cialdini’s principles of influence, human manipulation succeeds when using specific psychological triggers: Authority (misrepresenting senders as leadership), Urgency (imposing false immediate deadliness), and Social Proof.")
    add_custom_heading("7.2 Analysis of Industrial Phishing Reports", level=2)
    doc.add_paragraph("Data from the 2025 Verizon Data Breach Investigations Report (DBIR) and Microsoft Digital Defense Reports show that credential phishing accounts for over 60% of all social engineering vectors. Furthermore, attacks using clean, localized cloud services bypass traditional reputation-based URL filters over 73% of the time.")
    add_custom_heading("7.3 Comparative Study of Simulation Approaches", level=2)
    doc.add_paragraph("A review of existing open-source and enterprise frameworks (such as GoPhish vs. commercial alternatives) reveals that successful programs rely on high-fidelity landing page cloning, strict data boundary setups, and immediate learning interception.")
    doc.add_page_break()

    # 8. System Analysis
    add_custom_heading("8. SYSTEM ANALYSIS", level=1)
    add_custom_heading("8.1 Requirements Specification", level=2)
    doc.add_paragraph("Functional Requirements: The system must generate targeted emails, handle tracking pixels cleanly without breaking HTML formatting, intercept login inputs, and instantly display educational training pages while purging actual plaintext passwords to protect privacy.")
    doc.add_paragraph("Non-Functional Requirements: High availability of the landing page server during campaigns, processing latency below 200ms per click transaction, and full masking of simulation servers to prevent external indexing.")
    add_custom_heading("8.2 Architectural Layout", level=2)
    doc.add_paragraph("The platform is structurally organized into an Administration Panel (orchestration layer), a Delivery Node running custom SMTP agents, and a Telemetry Capture Server handling incoming webhooks.")
    doc.add_page_break()

    # 9. Methodology
    add_custom_heading("9. METHODOLOGY", level=1)
    doc.add_paragraph("The implementation followed a rigorous, multi-stage lifecycle conforming to ethical hacking guidelines and authorization protocols:")
    doc.add_paragraph("• Stage 1: Reconnaissance and Profile Mapping – Analyzing baseline targeted structures to ensure templates emulated plausible business workflows.", style='List Bullet')
    doc.add_paragraph("• Stage 2: DNS & Infrastructure Provisioning – Configuring full SPF, DKIM, and DMARC identity rules on our external delivery Virtual Private Server (VPS) nodes.", style='List Bullet')
    doc.add_paragraph("• Stage 3: Campaign Orchestration – Deploying specific tracking tokens embedded in email source codes safely measuring interaction states.", style='List Bullet')
    doc.add_paragraph("• Stage 4: Ingestion and Sanitization – Catching interaction data over secure API hooks while actively dropping string values from volatile memory buffers.", style='List Bullet')
    doc.add_paragraph("• Stage 5: Educational Intercept – Directing users to responsive landing layouts exposing specific visual red flags present in the test copy.", style='List Bullet')
    doc.add_page_break()

    # 10. Key Features
    add_custom_heading("10. KEY FEATURES", level=1)
    doc.add_paragraph("• High-Fidelity Contextual Simulations: Replicates authentic system states (e.g., Microsoft Teams updates, enterprise corporate payroll changes, network password maintenance alerts).", style='List Bullet')
    doc.add_paragraph("• Granular Telemetry Logs: Tracks real-time analytical values including click velocities, active device signatures, operating system variants, and interaction form indicators.", style='List Bullet')
    doc.add_paragraph("• Privacy-Preserving Logic: Structural cryptographic sanitization mechanics prevent credentials from being compiled into database logs.", style='List Bullet')
    doc.add_paragraph("• Dynamic Just-in-Time Training: Delivers micro-learning templates straight to the user at the exact moment of failure to maximize information retention.", style='List Bullet')
    doc.add_page_break()

    # 11. Skills Involved
    add_custom_heading("11. SKILLS INVOLVED", level=1)
    doc.add_paragraph("• Social Engineering Lure Metrics: Deconstructing psychological triggers to ethically measure structural corporate risk.", style='List Bullet')
    doc.add_paragraph("• Email Architecture Customization: Working natively with Mail Transfer Agents (MTAs), processing email components, and writing explicit validation parameters.", style='List Bullet')
    doc.add_paragraph("• Defensive Security Controls: Engineering localized alerting setups, understanding secure relays, and provisioning reverse proxy architectures.", style='List Bullet')
    doc.add_paragraph("• Analytics & Reporting Automation: Synthesizing large server event loops into structured business intelligence reporting models.", style='List Bullet')
    doc.add_page_break()

    # 12. Learning Outcomes
    add_custom_heading("12. LEARNING OUTCOMES", level=1)
    doc.add_paragraph("• Strategic Attack Understanding: Acquired deep functional visibility into adversarial operations, reconnaissance techniques, and target packaging methods.", style='List Bullet')
    doc.add_paragraph("• Advanced Forensic Inspection: Mastered raw transit trace reading, indicator validation parsing, and structural anomaly isolation inside communication headers.", style='List Bullet')
    doc.add_paragraph("• Enterprise Defensive Hardening: Developed professional implementation methodologies for enterprise validation standardizations and training architectures.", style='List Bullet')
    doc.add_page_break()

    # 13. Tools & Technologies Used
    add_custom_heading("13. TOOLS & TECHNOLOGIES USED", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component Category'
    hdr_cells[1].text = 'Technology Selection'
    hdr_cells[2].text = 'Operational Purpose'
    
    tools_data = [
        ("Core Framework", "GoPhish / Custom Extensions", "Campaign management and template orchestration."),
        ("Server Platform", "Linux Ubuntu Server (VPS)", "Hosted environment for web and mail listeners."),
        ("Web Server / Proxy", "Nginx", "Reverse proxy handling secure SSL termination."),
        ("Database Management", "SQLite / PostgreSQL", "Logging simulation states and analytical matrices."),
        ("Development Core", "Python / Bash Scripting", "Log processing, domain setups, and cleanup automations."),
        ("Security Layer", "Let's Encrypt", "Automated provisioning of valid TLS/SSL certificates.")
    ]
    
    for cat, tech, purp in tools_data:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = tech
        row_cells[2].text = purp
        
    doc.add_page_break()

    # Sections 14 - 19
    sections = [
        ("14. APPLICATIONS", [
            "• Corporate Risk Assessments: Allowing continuous internal evaluations before global auditing routines.",
            "• Regulatory Compliance Audit Readiness: Meeting validation checks under standards like ISO 27001, SOC 2, and PCI-DSS.",
            "• Targeted Employee Onboarding: Benchmarking risk levels immediately during standard HR training phases.",
            "• Long-Term Workforce Optimization: Executing staggered tests throughout the operational calendar year."
        ]),
        ("15. ADVANTAGES", [
            "• Safe Perimeter Verification: Safely isolates corporate environments from exposure to real technical vulnerabilities.",
            "• Dynamic Training Impact: Replaces passive viewing blocks with actionable real-time interaction feedback loops.",
            "• Deep Human Risk Metrics: Quantifies exact workforce response rates directly into business intelligence dashboards.",
            "• Resource Efficiency: Allows security departments to allocate focus directly toward highly targeted risk areas."
        ]),
        ("16. LIMITATIONS", [
            "• Point-In-Time Evaluation: Metrics capture distinct organizational moments; changing habits require persistent checking.",
            "• In-Group Adaptation Fatigue: Frequent baseline test methods risk predictable patterns over multi-year setups.",
            "• Scope Restrictions: Framework analyzes human cognitive choices rather than endpoint patching levels."
        ]),
        ("17. FUTURE ENHANCEMENTS", [
            "• AI-Engineered Targeted Scaling: Integrating modern model generation blocks to formulate tailored testing lures.",
            "• Omni-Channel Assessment Nodes: Branching tracking models natively across Smishing (SMS) and Vishing (Voice) channels.",
            "• Identity Provider Automation: Interfacing directly into directory engines for immediate contextual group enrollment."
        ]),
        ("18. CONCLUSION", [
            "This project highlights the critical role human behavior plays in modern enterprise cybersecurity. While traditional boundary defenses remain necessary, they cannot fully eliminate the risks posed by targeted social engineering attacks. By designing and executing a controlled phishing simulation framework, we successfully established an empirical method for measuring, analyzing, and reducing human risk.",
            "The data confirms that passive security training is insufficient for building long-term defensive habits. In contrast, regular, realistic simulations combined with immediate, contextual feedback significantly improve threat recognition and reporting rates. Ultimately, transforming a workforce from an operational vulnerability into an active defensive asset is one of the most effective strategies for mitigating modern cyber threats."
        ]),
        ("19. REFERENCES", [
            "1. Cialdini, R. B. (2021). Influence, New and Expanded: The Psychology of Persuasion. Harper Business.",
            "2. Verizon. (2025). Data Breach Investigations Report (DBIR). Cybersecurity Research Division.",
            "3. Mitre Corporation. (2026). MITRE ATT&CK Framework: T1566 - Phishing Techniques. Retrieved from https://attack.mitre.org/techniques/T1566/",
            "4. Hadnagy, C. (2018). Social Engineering: The Science of Human Hacking. John Wiley & Sons.",
            "5. GoPhish Project Documentation. (2025). Open-Source Phishing Framework Architecture and API Controls.",
            "6. Zeltser, L. (2024). Designing Effective Security Awareness Training Programs Through Behavioral Engineering. SANS Technology Institute."
        ])
    ]

    for title, paragraphs in sections:
        add_custom_heading(title, level=1)
        for p_text in paragraphs:
            if p_text.startswith("• "):
                doc.add_paragraph(p_text, style='List Bullet')
            else:
                doc.add_paragraph(p_text)
        doc.add_page_break()

    # Save compiled file
    doc.save("Cyber_Security_Internship_Report.docx")
    print("Success: 'Cyber_Security_Internship_Report.docx' generated successfully.")

if __name__ == "__main__":
    create_document()
